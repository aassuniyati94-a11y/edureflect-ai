import streamlit as st
from google import genai
from docx import Document
from io import BytesIO
from datetime import datetime

# 1. Konfigurasi API Key Gemini
# GANTI TULISAN DI BAWAH DENGAN API KEY GRATIS DARI GOOGLE AI STUDIO
client = genai.Client(api_key="AQ.Ab8RN6IKYpQo1CueodFw-l2p0eDKufTyf8BSxIJgdfiE0DCo3w")

# --- ISI DATA DEFAULT SEKOLAH ANDA DI SINI ---
DEFAULT_SEKOLAH = "SMP Negeri 1 Jakarta"
DEFAULT_KASEK = "Dr. H. Ahmad Sunarya, M.Pd."
DEFAULT_NIP_KASEK = "197508122000031002"
DEFAULT_GURU = "Budi Setiawan, S.Pd."
DEFAULT_NIP_GURU = "198810252015041003"

# --- PENGATURAN HALAMAN ---
st.set_page_config(page_title="EduReflect-AI", page_icon="🧬", layout="wide")

st.title("🧬 EduReflect-AI: Generator Modul Ajar Fase D Terintegrasi")
st.caption("Platform Penyelarasan Kurikulum Merdeka - Autentikasi Instan & Siap Cetak")

# Menambahkan catatan landasan pedagogis di sidebar
with st.sidebar:
    st.header("🔬 Landasan Pedagogis")
    st.markdown("""
    Aplikasi ini dirancang khusus untuk **Fase D (SMP)** dengan mematuhi dua pilar utama:
    1. **Dual-Aspect Reflection:** Pemisahan refleksi materi (kegiatan akhir) dan refleksi proses/emosional (penutup).
    2. **Trilogi JMM:** Integrasi aktivitas *Joyful, Mindful,* dan *Meaningful* di seluruh rangkaian pembelajaran.
    """)

# 2. Form Input
st.subheader("📋 Identitas Sekolah & Kurikulum")
col1, col2 = st.columns(2)
with col1:
    nama_sekolah = st.text_input("Nama Sekolah", value=DEFAULT_SEKOLAH)
    kelas = st.selectbox("Kelas", ["Kelas VII (7)", "Kelas VIII (8)", "Kelas IX (9)"])
    mata_pelajaran = st.text_input("Mata Pelajaran", placeholder="Contoh: IPA / Matematika")
    # TAH-DAH! INPUT MATERI POKOK SUDAH HADIR DI SINI:
    materi_pokok = st.text_input("Materi Pokok", placeholder="Contoh: Sistem Pencernaan / Persamaan Linear")
with col2:
    elemen = st.text_input("Elemen Pembelajaran", placeholder="Contoh: Aljabar / Ekologi")
    jumlah_jp = st.number_input("Jumlah Jam Pelajaran (JP)", min_value=1, max_value=10, value=2, step=1)
    jumlah_pertemuan = st.number_input("Untuk Berapa Pertemuan?", min_value=1, max_value=5, value=1, step=1)

st.subheader("✍️ Otentikasi & Tanda Tangan (Otomatis)")
col3, col4 = st.columns(2)
with col3:
    nama_kasek = st.text_input("Nama Kepala Sekolah", value=DEFAULT_KASEK)
    nip_kasek = st.text_input("NIP Kepala Sekolah", value=DEFAULT_NIP_KASEK)
with col4:
    nama_guru = st.text_input("Nama Guru Mata Pelajaran", value=DEFAULT_GURU)
    nip_guru = st.text_input("NIP Guru Mata Pelajaran", value=DEFAULT_NIP_GURU)

st.subheader("🎯 Capaian & Tujuan")
cp = st.text_area("Capaian Pembelajaran (CP)")
atp = st.text_area("Tujuan Pembelajaran (TP)")

def buat_file_word(teks_rpp):
    doc = Document()
    doc.add_heading('RENCANA PELAKSANAAN PEMBELAJARAN (RPP) / MODUL AJAR', level=1)
    for baris in teks_rpp.split('\n'):
        doc.add_paragraph(baris)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# 3. Eksekusi Tombol
if st.button("Susun RPP & Siapkan File Word"):
    if not cp or not atp or not mata_pelajaran or not materi_pokok:
        st.error("Mohon lengkapi data Mata Pelajaran, Materi Pokok, CP, dan TP!")
    else:
        with st.spinner("Gemini AI sedang menyusun modul lengkap..."):
            tanggal_sekarang = datetime.now().strftime("%d %B %Y")
            prompt_instruksi = f"""
            Anda adalah pakar Kurikulum Merdeka SMP (Fase D). Buatlah RPP/Modul Ajar resmi tingkat SMP dengan ketentuan wajib berikut:
            
            DATA MODUL:
            - Sekolah: {nama_sekolah}
            - Tingkat: SMP {kelas}
            - Mata Pelajaran: {mata_pelajaran}
            - Materi Pokok: {materi_pokok}
            - Elemen: {elemen}
            - Alokasi Waktu: {jumlah_jp} JP
            - Total Distribusi: {jumlah_pertemuan} Pertemuan
            - CP: {cp}
            - TP: {atp}
            
            DATA TANDA TANGAN:
            - Kepala Sekolah: {nama_kasek} (NIP: {nip_kasek})
            - Guru Mapel: {nama_guru} (NIP: {nip_guru})
            
            WAJIB MENGANDUNG PRINSIP TRILOGI JMM:
            Sepanjang RPP (dari Awal sampai Akhir), integrasikan secara nyata pendekatan untuk materi '{materi_pokok}':
            1. Joyful Learning, 2. Mindful Learning, 3. Meaningful Learning.
            Berikan tanda/keterangan tulisan seperti '[Joyful]' atau '[Mindful]' atau '[Meaningful]' di setiap aktivitas yang merepresentasikannya.
            
            PANDUAN DISTRIBUSI WAKTU (PENTING):
            Bagi langkah pembelajaran menjadi {jumlah_pertemuan} pertemuan secara logis dan runtut membahas materi '{materi_pokok}'. Setiap pertemuan harus memiliki struktur:
            1. KEGIATAN AWAL (Akomodasi unsur Joyful/Mindful)
            2. KEGIATAN INTI (Pembelajaran aktif dan bermakna)
            3. KEGIATAN AKHIR (Kesimpulan & [WAJIB] REFLEKSI PEMBELAJARAN/MATERI)
            4. PENUTUP (Tindak lanjut & [WAJIB] REFLEKSI PROSES/EMOSIONAL)
               
            Di akhir seluruh pertemuan, sertakan bagian 5. ASESMEN.
            
            BAGIAN PALING AKHIR (KOLOM TANDA TANGAN RESMI):
            Tuliskan kolom tanda tangan di bawah ini secara presisi:
            
            Mengetahui,
            Kepala Sekolah {nama_sekolah}                     Jakarta, {tanggal_sekarang}
                                                             Guru Mata Pelajaran
            
            
            ({nama_kasek})                                   ({nama_guru})
            NIP. {nip_kasek}                                 NIP. {nip_guru}
            """
            try:
                response = client.models.generate_content(model='gemini-2.5-flash', contents=prompt_instruksi)
                rpp_hasil = response.text
                st.success("✨ RPP Resmi Selesai Dibuat!")
                st.markdown("---")
                st.markdown(rpp_hasil)
                
                file_word = buat_file_word(rpp_hasil)
                st.download_button(
                    label="📥 Download RPP (.docx)",
                    data=file_word,
                    file_name=f"RPP_{mata_pelajaran}_{kelas.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Terjadi error pada Gemini API: {e}")
